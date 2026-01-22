from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def create_tabular_feasibility_doc():
    document = Document()

    # Title
    title = document.add_heading('PROJECT FEASIBILITY STUDY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('INTERNATIONAL JOB RECRUITMENT PORTAL')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Title'
    
    document.add_paragraph() # Spacer

    # Create the Main Table (1 Header row + 4 Content rows)
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.autofit = False 
    
    # Set Column Widths
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(1.5)

    # --- HEADER ROW ---
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feasibility Criteria'
    hdr_cells[1].text = 'Weight (Percentage)'
    
    # Style Header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # --- ROW 1: OPERATIONAL FEASIBILITY ---
    row1 = table.rows[1].cells
    row1[1].text = '30 %'
    row1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row1[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    c1 = row1[0]
    p = c1.paragraphs[0]
    p.add_run('6.1.1 Operational Feasibility:').bold = True
    c1.add_paragraph('This system is highly operationally feasible as it directly addresses the critical need for a centralized, transparent platform for international recruitment.')
    
    c1.add_paragraph('Candidate Benefit:').bold = True
    c1.add_paragraph('Candidates gain a unified interface to upload documents (Passport, PCC, Resume) and track application status in real-time. The dashboard eliminates manual follow-ups by providing instant status visibility.')
    
    c1.add_paragraph('Organization Benefit:').bold = True
    c1.add_paragraph('The "3-User Architecture" (Admin/Recruiter/Candidate) streamlines the vetting process. Automated role-based routing ensures smooth data flow without manual intervention.')

    # --- ROW 2: TECHNICAL FEASIBILITY ---
    row2 = table.rows[2].cells
    row2[1].text = '30 %'
    row2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row2[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    c2 = row2[0]
    p = c2.paragraphs[0]
    p.add_run('6.1.2 Technical Feasibility:').bold = True
    c2.add_paragraph('The project relies on the robust, open-source MERN stack, ensuring high scalability and control:')
    
    c2.add_paragraph('Frontend Maturity:').bold = True
    c2.add_paragraph('• React + JavaScript (ES6+): Ensures rapid development with a flexible component model.')
    c2.add_paragraph('• Vite: Provides faster build times than traditional bundlers.')
    c2.add_paragraph('• Tailwind CSS: Enables rapid, responsive UI development.')
    
    c2.add_paragraph('Backend & Database Reliability:').bold = True
    c2.add_paragraph('• Node.js & Express: Lightweight, non-blocking I/O ideal for handling concurrent requests.')
    c2.add_paragraph('• MongoDB Atlas: Scalable NoSQL cloud database for flexible data modeling.')
    c2.add_paragraph('• GridFS Storage: Secure, integrated file storage for handling large resumes and documents directly in the database.')

    # --- ROW 3: ECONOMIC FEASIBILITY ---
    row3 = table.rows[3].cells
    row3[1].text = '20 %'
    row3[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row3[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    c3 = row3[0]
    p = c3.paragraphs[0]
    p.add_run('6.1.3 Economic Feasibility:').bold = True
    c3.add_paragraph('Cost to Develop: Development costs are minimal as the entire stack is open-source.')
    
    c3.add_paragraph('Infrastructure Savings:').bold = True
    c3.add_paragraph('• Fullstack JavaScript: Reduces resource costs by unifying Frontend and Backend languages.')
    c3.add_paragraph('• Free Tier Hosting: MongoDB Atlas (Database) and Vercel/Render (App Hosting) offer generous free tiers.')
    
    c3.add_paragraph('ROI: The platform automates manual agency tasks, reducing labor costs and turnover time.')

    # --- ROW 4: SCHEDULE FEASIBILITY ---
    row4 = table.rows[4].cells
    row4[1].text = '20 %'
    row4[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row4[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    c4 = row4[0]
    p = c4.paragraphs[0]
    p.add_run('6.1.4 Schedule Feasibility:').bold = True
    c4.add_paragraph('The project is feasible within the 8-week timeline due to parallel task allocation.')
    
    c4.add_paragraph('Timeline Overview:').bold = True
    # Nested table for timeline
    t_inner = c4.add_table(rows=5, cols=2)
    t_inner.style = 'Table Grid'
    t_inner.cell(0,0).text = 'Phase'
    t_inner.cell(0,1).text = 'Duration'
    t_inner.cell(1,0).text = 'Requirements & Design'
    t_inner.cell(1,1).text = 'Week 1-2'
    t_inner.cell(2,0).text = 'Development (MERN Stack)'
    t_inner.cell(2,1).text = 'Week 3-5'
    t_inner.cell(3,0).text = 'Integration & Testing'
    t_inner.cell(3,1).text = 'Week 6-7'
    t_inner.cell(4,0).text = 'Deployment'
    t_inner.cell(4,1).text = 'Week 8'
    
    c4.add_paragraph()
    c4.add_paragraph('Risk Mitigation:').bold = True
    c4.add_paragraph('• Unified Language: Using JavaScript across the stack speeds up development.')
    c4.add_paragraph('• Libraries: Using proven libraries like Mongoose and Multer reduces custom coding effort.')

    document.save('Project_Feasibility_Study.docx')

if __name__ == "__main__":
    create_tabular_feasibility_doc()
