from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE MARGINS ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── HELPER FUNCTIONS ──────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)  # teal
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # dark navy
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x13, 0x4E, 0x4A)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def kv(key, value):
    """Bold key + normal value on one line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    k = p.add_run(f"{key}: ")
    k.bold = True
    k.font.size = Pt(11)
    v = p.add_run(value)
    v.font.size = Pt(11)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0D9488')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(60)
tr = title_p.add_run("GlobalAKJobs — Maldives Career Platform")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Project Technical Documentation")
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("Date: March 2026")
dr.font.size = Pt(11)
dr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# Page break before content
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1("1. Project Overview")
divider()
body(
    "GlobalAKJobs is a full-stack web application developed as a job portal specifically "
    "designed for the Maldives employment market. The platform connects job seekers (Candidates) "
    "with recruitment agencies (Agents) and is managed by platform Administrators. It supports "
    "the complete lifecycle of job seeking — from browsing and applying to jobs, uploading "
    "application documents, tracking application status, to managing recruiter and administrative "
    "operations via dedicated dashboards."
)
body(
    "The platform's name 'GlobalAKJobs' with the tagline 'Island Jobs Simplified' reflects its "
    "goal of making Maldivian employment opportunities accessible and streamlined for both local "
    "and international candidates."
)

heading2("Key Features")
bullet("Public job browsing and search with category and location filters")
bullet("Candidate registration, login (email/password and Google OAuth), and profile management")
bullet("Job application submission with document uploads (Resume, Identity, Certificates, PCC, Good Standing)")
bullet("Agent (Recruiter) dashboard for managing job listings and reviewing applicants")
bullet("Admin dashboard for approving agents, managing users, monitoring the platform")
bullet("Email-based password reset with time-limited secure tokens")
bullet("In-app notification system for agents and candidates")
bullet("Saved jobs feature for candidates")
bullet("Role-based access control (Candidate / Agent / Admin)")

# ═══════════════════════════════════════════════════════════════════════════════
#  2. REQUIREMENTS SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("2. Requirements Specifications")
divider()

# ─── 2.1 HARDWARE REQUIREMENTS ───────────────────────────────────────────────
heading2("2.1  Hardware Requirements")
body(
    "The application is a web-based platform accessible through any modern browser. "
    "The following hardware specifications are recommended for development, deployment servers, "
    "and end users."
)

heading3("Development Machine (Minimum)")
kv("Processor", "Intel Core i5 (8th Gen) / AMD Ryzen 5 or equivalent — 2.0 GHz or faster")
kv("RAM", "8 GB (16 GB recommended for running both frontend and backend simultaneously)")
kv("Storage", "256 GB SSD (at least 10 GB free for project files, Node modules, and build artifacts)")
kv("Operating System", "Windows 10/11, macOS 10.15+, or Ubuntu 20.04+")
kv("Network", "Stable broadband internet connection (minimum 10 Mbps) for API calls, npm installs, and cloud DB access")
kv("Display", "1280 × 720 resolution minimum; 1920 × 1080 recommended for comfortable development")

heading3("Development Machine (Recommended)")
kv("Processor", "Intel Core i7/i9 or AMD Ryzen 7 — 3.0 GHz+")
kv("RAM", "16 GB or more")
kv("Storage", "512 GB NVMe SSD")
kv("Network", "50 Mbps+ broadband")

heading3("Production Server — Backend (Render.com Free/Starter Tier)")
kv("CPU", "Shared vCPU (0.1 CPU on free tier; 0.5–1 vCPU on starter)")
kv("RAM", "512 MB (free tier) — 512 MB to 2 GB (starter tier)")
kv("Storage", "Ephemeral disk (no persistent disk on free tier; files stored in MongoDB as Base64)")
kv("Network", "Auto-scaled; HTTPS enforced via Render's built-in TLS termination")

heading3("Production Server — Frontend (Netlify / Vercel CDN)")
kv("Infrastructure", "Global CDN — no dedicated hardware; static files served from edge nodes worldwide")
kv("Bandwidth", "100 GB/month (free tier); unlimited on paid plans")
kv("Build Memory", "3 GB build container provided by Netlify/Vercel")

heading3("Database Server — MongoDB Atlas (Free Tier / M0)")
kv("Storage", "512 MB (M0 Sandbox — shared cluster)")
kv("RAM", "Shared RAM (Atlas M0)")
kv("vCPU", "Shared compute")
kv("Network", "Auto-managed; supports TLS connections from Render backend")
kv("Cluster Region", "Cloud cluster on MongoDB Atlas — globalaKjobs cluster (AWS)")

heading3("End-User Devices (Client Requirements)")
body("Since the application is a progressive web application served through a browser:")
kv("Device", "Desktop, laptop, tablet, or smartphone")
kv("Screen Resolution", "Minimum 360 × 640 (mobile); 1280 × 720 (desktop)")
kv("RAM", "2 GB minimum (4 GB recommended for smooth browser performance)")
kv("Storage", "No local storage required; browser cache (~5 MB)")
kv("Browser", "Google Chrome 100+, Mozilla Firefox 100+, Microsoft Edge 100+, Safari 15+")
kv("Network", "Minimum 2 Mbps internet connection (4G/LTE or broadband)")

# ─── 2.2 SOFTWARE REQUIREMENTS ───────────────────────────────────────────────
heading2("2.2  Software Requirements")
body(
    "The following software components are required to develop, run, and deploy the GlobalAKJobs platform."
)

heading3("Operating System")
kv("Development", "Windows 10/11, macOS, or Linux (Ubuntu 20.04+)")
kv("Production", "Linux (Render servers use Debian-based Linux containers)")

heading3("Runtime Environments")
kv("Node.js", "v18.x LTS or later (required for both frontend build tooling and backend server)")
kv("NPM", "v9.x or later (comes bundled with Node.js)")
kv("Python (optional)", "3.x — used only for utility/seeding scripts if needed")

heading3("Version Control")
kv("Git", "v2.x or later")
kv("GitHub", "Cloud-hosted repository for source code, collaboration, and CI/CD integration")

heading3("Development Tools / IDEs")
kv("Primary IDE", "Visual Studio Code (VS Code) with the following extensions:")
bullet("ESLint — JavaScript linting and code quality", level=1)
bullet("Prettier — Code formatting", level=1)
bullet("Tailwind CSS IntelliSense — Autocomplete for Tailwind classes", level=1)
bullet("MongoDB for VS Code — MongoDB database explorer", level=1)
bullet("GitLens — Advanced Git integration", level=1)
kv("API Testing Tool", "Postman or Thunder Client (VS Code extension) — for testing REST API endpoints")
kv("Browser DevTools", "Chrome DevTools — for debugging, network inspection, and performance profiling")

heading3("Package Managers")
kv("NPM", "Used for managing both frontend and backend JavaScript dependencies")

heading3("Database Software")
kv("MongoDB Atlas", "Cloud-hosted NoSQL database (MongoDB v6.0+) — used in production")
kv("MongoDB Compass (optional)", "GUI tool for local database inspection and management")
kv("Mongoose ODM", "v8.0.3 — Object Data Modelling library for MongoDB, used in the backend")

heading3("Environment Configuration")
kv(".env files", "Separate .env files for backend and frontend storing sensitive keys (never committed to git)")
body("Backend .env variables:")
bullet("MONGODB_URI — MongoDB Atlas connection string", level=1)
bullet("GOOGLE_CLIENT_ID — Google OAuth 2.0 Client ID from Google Cloud Console", level=1)
bullet("GOOGLE_CLIENT_SECRET — Google OAuth 2.0 Client Secret", level=1)
bullet("GOOGLE_CALLBACK_URL — OAuth callback URL", level=1)
bullet("GOOGLE_ALLOWED_ORIGINS — Comma-separated list of allowed redirect origins", level=1)
bullet("EMAIL_USER — Gmail address used for Nodemailer (dev fallback)", level=1)
bullet("EMAIL_APP_PASSWORD — Gmail App Password for Nodemailer (dev fallback)", level=1)
bullet("RESEND_API_KEY — Resend.com API key used in production for email sending", level=1)
bullet("ALLOWED_ORIGINS — Comma-separated list of allowed CORS frontend origins", level=1)
body("Frontend .env variables:")
bullet("VITE_GOOGLE_CLIENT_ID — Google OAuth Client ID (must match backend)", level=1)
bullet("VITE_API_URL — Backend API base URL (http://localhost:5000 for dev; Render URL for prod)", level=1)

heading3("Security and Authentication")
kv("Google OAuth 2.0", "Authentication via Google Identity Services (GIS) — Google Cloud Console project required")
kv("JWT-like token validation", "Google ID tokens verified via Google's tokeninfo endpoint")
kv("Password Reset Tokens", "Node.js crypto module — 32-byte random hex tokens with 15-minute expiry")
kv("CORS Policy", "Configured via the cors npm package — only whitelisted origins allowed")
kv("Helmet.js", "v8.1.0 — HTTP security headers middleware in Express")
kv("Rate Limiting", "express-rate-limit v8.2.1 — prevents brute-force attacks on API endpoints")

heading3("Email Services")
kv("Production Email", "Resend.com HTTP API — used for sending password reset emails on Render (bypasses SMTP port restrictions)")
kv("Development Email", "Nodemailer with Ethereal.email — captures emails locally without sending real emails")

# ═══════════════════════════════════════════════════════════════════════════════
#  3. DEVELOPMENT TOOLS / TECHNOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("3. Development Tools / Technology")
divider()

# ─── 3.1 FRONTEND ────────────────────────────────────────────────────────────
heading2("3.1  Frontend")
body(
    "The frontend is a modern Single Page Application (SPA) built using React 19 with Vite 6 "
    "as the build tool and development server. It is styled using Tailwind CSS and enhanced "
    "with smooth animations via Framer Motion."
)

heading3("Core Framework & Library")
kv("React", "v19.2.3 — UI component library by Meta (Facebook). Used for building all user interface components, managing state, and handling user interactions through a virtual DOM.")
kv("React DOM", "v19.2.3 — Renderer that connects React components to the real browser DOM.")

heading3("Build Tool")
kv("Vite", "v6.2.0 — Next-generation frontend build tool. Provides instant Hot Module Replacement (HMR) during development and optimised production builds using Rollup. Configured in vite.config.js.")
kv("@vitejs/plugin-react", "v5.0.0 — Official Vite plugin for React; enables JSX transformation and fast refresh.")

heading3("Routing")
kv("React Router DOM", "v7.11.0 — Client-side routing library. Used to define all application routes (/, /jobs, /job/:id, /login, /register, /dashboard, /recruiter, /admin, /profile, /saved-jobs, /reset-password, etc.) and handle protected route access by role.")

heading3("Styling")
kv("Tailwind CSS", "v3.4.19 — Utility-first CSS framework. Provides all layout, spacing, colour, and responsive design utilities used throughout the application without writing custom CSS.")
kv("PostCSS", "v8.5.6 — CSS transformation tool (required by Tailwind CSS).")
kv("Autoprefixer", "v10.4.27 — PostCSS plugin that automatically adds vendor prefixes for cross-browser compatibility.")

heading3("Animation")
kv("Framer Motion", "v12.34.3 — Production-ready motion library for React. Used for page transitions, card animations, dropdown animations, and other micro-interactions throughout the UI.")

heading3("Icons")
kv("Lucide React", "v0.562.0 — Icon library providing clean, consistent SVG icons used throughout the application (search, briefcase, location, user, bell icons, etc.).")

heading3("Backend-as-a-Service / Supabase Integration")
kv("@supabase/supabase-js", "v2.90.1 — JavaScript client for Supabase. Integrated for supplementary backend-as-a-service features (authentication helpers, real-time subscriptions, or storage as needed).")
kv("Supabase CLI", "v2.72.7 (devDependency) — Used for local development and database schema management with Supabase.")

heading3("TypeScript Support")
kv("TypeScript", "v5.8.2 (devDependency) — Type definitions available; the project uses JSX (JavaScript) for component files but TypeScript is configured for type safety via @types/node.")
kv("@types/node", "v22.14.0 — Node.js type definitions for TypeScript tooling compatibility.")

heading3("Authentication on Frontend")
kv("Google Identity Services (GIS)", "The Google Sign-In button is rendered using the official Google Identity Services library loaded via script tag. The user's Google credential (JWT ID token) is captured and sent to the backend /api/auth/google endpoint for server-side verification.")
kv("AuthContext", "A global React context (AuthContext.jsx) manages the authenticated user's session, persists login state in localStorage (excluding avatar for quota reasons), and provides login, loginWithGoogle, logout, and updateUser functions to all components.")

heading3("Key Frontend Pages & Components")
body("The frontend consists of the following major pages:")
kv("HomePage", "Landing page with hero section, featured jobs, and call-to-action")
kv("BrowseJobsPage", "Full job listing with search, category filter, and location filter")
kv("JobDetailPage", "Individual job detail view with application form and document upload")
kv("CandidateLoginPage", "Login and registration page for candidates (email/password + Google OAuth)")
kv("AgentLoginPage", "Login page specifically for recruitment agents")
kv("AdminLoginPage", "Secure admin login page")
kv("CandidateDashboard", "Candidate's personal dashboard showing applications and status")
kv("RecruiterDashboard", "Agent's portal for managing job posts and reviewing applicants")
kv("AdminDashboard", "Full platform management panel for administrators")
kv("ProfilePage", "User profile management (name, contact, location, skills, avatar, password change)")
kv("SavedJobsPage", "List of jobs saved/bookmarked by the candidate")
kv("ResetPasswordPage", "Password reset page accessed via emailed token link")
kv("AgentRegistrationPage", "Agent self-registration form (pending admin approval)")
kv("SupportPage", "Contact/support page")
kv("SuccessPage", "Confirmation page shown after successful application submission")

heading3("Frontend Port & Configuration")
kv("Development Port", "3001 (configured in vite.config.js)")
kv("API Base URL (dev)", "http://localhost:5000 (set in VITE_API_URL env variable)")
kv("API Base URL (prod)", "Render backend deployment URL (set in Netlify/Vercel environment settings)")

# ─── 3.2 BACKEND ─────────────────────────────────────────────────────────────
doc.add_page_break()
heading2("3.2  Backend")
body(
    "The backend is a RESTful API server built with Node.js and Express.js, connected to "
    "a cloud-hosted MongoDB Atlas database via Mongoose ODM. It handles all business logic, "
    "authentication, file storage, and email services."
)

heading3("Runtime")
kv("Node.js", "v18.x LTS — JavaScript runtime for executing server-side code. Uses ES Modules (type: module in package.json).")

heading3("Web Framework")
kv("Express.js", "v4.18.2 — Minimal and flexible Node.js web application framework used to define all REST API routes, middleware, and request/response handling.")

heading3("Database ODM")
kv("Mongoose", "v8.0.3 — MongoDB Object Data Modelling (ODM) library. Provides schema definitions, data validation, model methods, and query building for all MongoDB collections.")

heading3("Database Models (Collections)")
body("The following Mongoose models define the database schema:")
kv("Profile", "Stores all users (Candidates, Agents, Admins). Fields: id (UUID), role, full_name, email, password, contact_number, agency_name, license_number, skills, avatar (Base64), experience_years, location, savedJobs, status, requiresPasswordChange, temporaryPassword, agencyId, googleId, resetPasswordToken, resetPasswordExpires, timestamps.")
kv("Job", "Stores job listings. Fields: id (UUID), title, company, location, category (Hospitality/Construction/Healthcare/IT/Education/Retail/Manufacturing/Tourism/Fishing/Agriculture/Other), salary_range, posted_date, status (OPEN/CLOSED), description, requirements, timestamps.")
kv("Application", "Stores job applications. Fields: id (UUID), job_id, candidate_name, email, contact_number, agent_id, resume (Base64), identity (Base64), certificates (Base64), pcc (Base64), goodStanding (Base64), status (PENDING/REVIEWING/ACCEPTED/REJECTED/APPROVED/HOLD/SELECTED), visibility_request_status, applied_at, timestamps.")
kv("Agency", "Stores recruitment agency information linked to Agent profiles.")
kv("JobRequest", "Stores job posting requests submitted by agents for admin approval.")
kv("Document", "Stores additional documents linked to applications.")
kv("Subscription", "Stores notification or alert subscription preferences.")
kv("Notification", "Stores in-app notifications for users; served via notification_routes.js.")

heading3("File Uploads")
kv("Multer", "v1.4.4 — Express middleware for handling multipart/form-data (file uploads). Configured to use memory storage (files temporarily held in memory buffer, then converted to Base64 and saved directly to MongoDB).")
kv("File Size Limit", "10 MB per file")
kv("Supported File Types", "PDFs, images (JPEG, PNG), and other document formats")
kv("Storage Method", "Base64-encoded data URLs stored directly in MongoDB fields — no separate file server required")

heading3("Authentication Middleware")
kv("CORS", "v2.8.5 — Cross-Origin Resource Sharing middleware. Configured to whitelist specific frontend origins (localhost:3000, localhost:5173, localhost:3001, plus production URLs via ALLOWED_ORIGINS env variable).")
kv("Helmet", "v8.1.0 — Sets secure HTTP response headers (X-Content-Type-Options, X-Frame-Options, Content-Security-Policy, etc.)")
kv("express-rate-limit", "v8.2.1 — Limits repeated API requests from the same IP to prevent abuse and brute-force attacks.")

heading3("Security Utilities")
kv("Node.js crypto (built-in)", "Used to generate cryptographically secure 32-byte random hex tokens for password reset functionality.")
kv("UUID (uuid v9.0.1)", "Generates unique v4 UUIDs for the custom id field on all database documents (Profile, Job, Application).")
kv("dotenv", "v16.3.1 — Loads environment variables from the .env file into process.env at runtime.")

heading3("Email Service")
kv("Nodemailer", "v8.0.1 — Node.js library for sending emails. Used as fallback in development with Ethereal.email test accounts that capture emails without sending them.")
kv("Resend HTTP API", "In production (on Render), SMTP outbound ports are blocked. The application uses Resend.com's HTTP API (HTTPS port 443) to reliably send password reset emails. The RESEND_API_KEY environment variable activates this path.")

heading3("API Routes Overview")
body("The backend exposes the following major REST API endpoint groups:")
kv("POST /api/auth/register", "Register a new Candidate or Agent account")
kv("POST /api/auth/login", "Login with email and password (role-specific)")
kv("POST /api/auth/google", "Google OAuth 2.0 login / auto-registration for candidates")
kv("POST /api/auth/forgot-password", "Send password reset email with secure time-limited token")
kv("GET /api/auth/validate-reset-token/:token", "Validate a password reset token")
kv("POST /api/auth/reset-password-token", "Apply new password using a valid reset token")
kv("PUT /api/auth/change-password", "First-time agent password change (admin-temp password flow)")
kv("PUT /api/auth/password", "Authenticated user password update from profile")
kv("GET /api/profile/:id", "Fetch a user's full profile")
kv("PUT /api/profile/:id", "Update profile fields (name, contact, location, skills, experience)")
kv("POST /api/profile/:id/avatar", "Upload/update profile photo (stored as Base64 in MongoDB)")
kv("POST /api/profile/:id/save-job", "Toggle saved/bookmarked job for a candidate")
kv("GET /api/profile/:id/saved-jobs", "Fetch all saved jobs for a candidate")
kv("GET /api/admin/pending-agents", "Fetch all agents awaiting admin approval")
kv("PUT /api/admin/agents/:id/approve", "Approve an agent account")
kv("GET/POST /api/jobs", "List all jobs / create a new job listing")
kv("GET/PUT/DELETE /api/jobs/:id", "Get, update, or delete a specific job")
kv("POST /api/applications", "Submit a job application with document uploads")
kv("GET /api/applications", "Fetch applications (filtered by role)")
kv("PUT /api/applications/:id/status", "Update application status (agent/admin)")
kv("/api/notifications/*", "Notification routes for in-app alerts (served via notification_routes.js)")
kv("GET /api/health", "Health check endpoint returning server and DB status")

heading3("Backend Port & Start Commands")
kv("Default Port", "5000 (configurable via PORT environment variable)")
kv("Production Start", "npm start  (runs: node server.js)")
kv("Development Start", "npm run dev  (runs: nodemon server.js — auto-restarts on file changes)")

# ─── 3.3 DEVELOPMENT TOOLS ───────────────────────────────────────────────────
doc.add_page_break()
heading2("3.3  Development Tools")
body("The following tools were used throughout the development process.")

heading3("Code Editor & IDE")
kv("Visual Studio Code (VS Code)", "Primary code editor. Supports JavaScript, JSX, TypeScript, JSON and has excellent extensions for React and Node.js development.")

heading3("Version Control")
kv("Git", "Distributed version control system used for tracking all changes to the codebase throughout development.")
kv("GitHub", "Cloud repository hosting platform used for team collaboration, code reviews, branch management, and as the deployment source for both Render (backend) and Netlify/Vercel (frontend).")

heading3("API Testing")
kv("Postman", "Industry-standard GUI tool used to manually test all backend REST API endpoints during development (testing auth, job CRUD, application submission, etc.).")

heading3("Database Management")
kv("MongoDB Compass", "Official MongoDB GUI tool used to visually inspect, query, and manage the MongoDB Atlas collections during development and debugging.")
kv("MongoDB Atlas Dashboard", "Cloud web interface for managing the cluster, monitoring queries, configuring access, and managing database users.")

heading3("Browser Developer Tools")
kv("Chrome DevTools", "Used for inspecting the React component tree (with React DevTools extension), debugging network requests, checking console logs, and profiling frontend performance.")

heading3("DNS & Network Tunnelling")
kv("ngrok", "Used during development to expose the local frontend (port 3001) to the internet — enabling webhook testing and mobile device testing. The allowed host lili-remorseless-cicely.ngrok-free.dev is configured in vite.config.js.")

heading3("Package Management")
kv("NPM (Node Package Manager)", "Used to install, update, and manage all Node.js dependencies for both the frontend and backend (separate node_modules directories in /frontend and /backend).")

heading3("Environment Management")
kv(".env Files", "Separate .env configuration files for the frontend (VITE_ prefixed variables) and backend are used to keep all secrets and environment-specific configuration out of the source code.")

# ─── 3.4 DEPLOYMENT ──────────────────────────────────────────────────────────
heading2("3.4  Deployment")
body(
    "The GlobalAKJobs application follows a split deployment architecture: "
    "the frontend (static SPA) and backend (Node.js server) are deployed independently "
    "on separate cloud platforms."
)

heading3("Backend Deployment — Render.com")
kv("Platform", "Render.com — Cloud platform for hosting backend web services")
kv("Service Type", "Web Service (Node.js)")
kv("Deployment Trigger", "Automatic deployment triggered by every push to the main branch of the GitHub repository")
kv("Build Command", "npm install")
kv("Start Command", "npm start  (which runs: node server.js)")
kv("Environment Variables on Render", "MONGODB_URI, GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, GOOGLE_CALLBACK_URL, GOOGLE_ALLOWED_ORIGINS, RESEND_API_KEY, ALLOWED_ORIGINS, EMAIL_USER, EMAIL_APP_PASSWORD — all set in the Render Dashboard > Environment tab")
kv("Port", "Render automatically assigns a PORT environment variable; the server uses process.env.PORT || 5000")
kv("TLS/HTTPS", "Render provides automatic TLS certificates — the backend is served over HTTPS")
kv("DNS Fix", "dns.setDefaultResultOrder('ipv4first') applied in server.js to overcome Render free tier's IPv6 outbound block (required for email sending and external API calls)")
kv("Limitations (Free Tier)", "Server spins down after 15 minutes of inactivity (cold-start delay on first request); 512 MB RAM; no persistent disk storage (all files stored in MongoDB)")

heading3("Frontend Deployment — Netlify / Vercel")
kv("Primary Platform", "Netlify (production) with Vercel also configured (vercel.json present in the frontend directory)")
kv("Deployment Trigger", "Automatic deployment on push to the main GitHub branch")
kv("Build Command", "npm run build  (Vite builds the production bundle into the /dist folder)")
kv("Publish Directory", "dist/")
kv("Environment Variables", "VITE_GOOGLE_CLIENT_ID and VITE_API_URL set in the Netlify/Vercel dashboard — these are inlined at build time by Vite")
kv("CDN", "Files are distributed globally via Netlify's Edge CDN — no server-side rendering; purely static SPA")
kv("Routing Fix", "vercel.json / Netlify redirects configured to route all paths (/*) to index.html enabling React Router client-side navigation to work correctly without 404 errors")
kv("HTTPS", "Automatic HTTPS/TLS with free SSL certificate")

heading3("Database Deployment — MongoDB Atlas")
kv("Service", "MongoDB Atlas (M0 Sandbox — free shared cluster)")
kv("Cluster Name", "GlobalAKjobs  (connection string cluster: globalakjobs.v9dyxlu.mongodb.net)")
kv("Database Name", "GlobalAKJobs-DB")
kv("Cloud Provider & Region", "AWS (region auto-selected by Atlas)")
kv("Access Control", "IP whitelisting (Render's outbound IPs whitelisted; 0.0.0.0/0 for development) and database user credentials in connection string")
kv("Connection String Format", "mongodb+srv://<user>:<password>@globalakjobs.v9dyxlu.mongodb.net/GlobalAKJobs-DB?retryWrites=true&w=majority&appName=GlobalAKjobs")

# ─── 3.5 HOSTING ─────────────────────────────────────────────────────────────
doc.add_page_break()
heading2("3.5  Hosting")
body(
    "The application uses a modern cloud-based hosting architecture with each component "
    "hosted on a purpose-built platform optimised for its role."
)

heading3("Frontend Hosting — Netlify")
kv("Provider", "Netlify (netlify.com)")
kv("Type", "Static Site Hosting with Global CDN")
kv("What is hosted", "The production build output of the React/Vite application (HTML, CSS, JavaScript bundles, images)")
kv("How it works", "After running npm run build, Vite compiles the entire React application into static files in the /dist directory. These files are uploaded to Netlify's global Content Delivery Network (CDN), which serves them from servers closest to the user's geographic location for maximum speed.")
kv("Custom Domain", "A custom domain can be configured via Netlify's domain management (free SSL included)")
kv("Build Triggers", "Connected to the GitHub repository; every push to the main branch triggers a new build and deployment automatically")
kv("Environment Variables", "VITE_API_URL and VITE_GOOGLE_CLIENT_ID are configured as build-time environment variables in the Netlify dashboard")
kv("SPA Support", "A redirect rule ensures all URL paths are served by index.html so React Router handles navigation without server-side route resolution")

heading3("Backend Hosting — Render.com")
kv("Provider", "Render (render.com)")
kv("Type", "Web Service — Node.js server hosting")
kv("What is hosted", "The Express.js REST API server (server.js) that handles all business logic, authentication, database operations, and file management")
kv("How it works", "Render pulls the backend code from the GitHub repository, runs npm install to install dependencies, then starts the server with node server.js. It manages the server process, restarts it if it crashes, and provides a public HTTPS URL for the API.")
kv("Free Tier Behaviour", "On the free tier, the server enters a sleep mode after 15 minutes of no traffic. The first request after sleep causes a cold-start delay of 30–60 seconds. Paid tiers keep the server always awake.")
kv("Persistent URL", "Render provides a stable .onrender.com URL for the backend (e.g., https://globalakjobs-backend.onrender.com)")
kv("Environment Configuration", "All environment variables (secrets, API keys, database URI) are set through the Render dashboard and injected into the server process at runtime")

heading3("Database Hosting — MongoDB Atlas")
kv("Provider", "MongoDB Atlas (cloud.mongodb.com)")
kv("Type", "Fully managed cloud NoSQL database service")
kv("What is hosted", "All application data: user profiles, job listings, job applications (with embedded Base64 documents), notifications, agencies, job requests")
kv("How it works", "The backend connects to MongoDB Atlas using a mongodb+srv:// connection string via Mongoose. Atlas handles all database infrastructure — hardware, backups, scaling, and security patching automatically.")
kv("Data Durability", "Atlas replicates data across multiple availability zones (replica set) ensuring data is not lost even if one server fails")
kv("Security", "Database access is restricted by IP whitelist and username/password credentials. All connections use TLS encryption.")

heading3("Authentication Hosting — Google Cloud Console")
kv("Provider", "Google Cloud (console.cloud.google.com)")
kv("Purpose", "Google OAuth 2.0 authentication service")
kv("What is configured", "An OAuth 2.0 Client ID and Client Secret are created in the Google Cloud Console. The Authorised JavaScript origins (frontend URL) and Authorised redirect URIs (backend callback URL) are configured to allow Google Sign-In to work.")
kv("How it works", "When a user clicks 'Sign in with Google', the Google Identity Services library opens a Google sign-in popup. After authentication, Google returns a JWT ID token (credential) to the frontend. The frontend sends this credential to the backend /api/auth/google endpoint, which verifies it against Google's tokeninfo API before creating or logging in the user.")

heading3("Email Service Hosting — Resend.com")
kv("Provider", "Resend (resend.com)")
kv("Purpose", "Transactional email delivery for password reset emails")
kv("How it works", "When a candidate requests a password reset, the backend calls Resend's HTTPS API (port 443) with the email content. Resend delivers the email to the user's inbox. This method is used in production because Render's free tier blocks outbound SMTP connections (port 587/465).")
kv("Development Alternative", "In development (without RESEND_API_KEY set), Nodemailer + Ethereal.email captures emails locally and logs a preview URL to the console — no real email is sent")

# ─── 3.6 ADDITIONAL TOOLS ────────────────────────────────────────────────────
heading2("3.6  Additional Tools & Technologies")

heading3("DNS over IPv4 Fix")
kv("dns.setDefaultResultOrder('ipv4first')", "Applied at the top of server.js. This Node.js native configuration forces all DNS lookups to prefer IPv4 addresses. This was necessary because Render's free tier blocks IPv6 outbound connections, which would prevent the backend from connecting to external services (MongoDB Atlas, Google APIs, Resend API) if IPv6 was tried first.")

heading3("Google Identity Services (GIS) Library")
kv("google.accounts.id", "Google's official JavaScript library loaded in the frontend for rendering the 'Sign in with Google' button and collecting the user's Google credential (JWT ID token) for server-side verification.")

heading3("lucide-react Icons")
kv("Lucide React", "Open-source icon library with over 1000 clean, consistent SVG icons. Provides icons for navigation, job categories, user profiles, notifications, search, and buttons throughout the application.")

heading3("Framer Motion Animations")
kv("Framer Motion", "Used for all UI animations including: page transitions, card hover/tap effects, modal open/close animations, sidebar slide-ins on dashboards, and staggered list animations for job cards.")

heading3("React Context API")
kv("AuthContext", "Manages global authentication state. Wraps the entire application so any component can access the current logged-in user's data and auth functions (login, loginWithGoogle, logout, updateUser) without prop-drilling.")
kv("PopupContext", "Manages global popup/modal state for showing notification popups and system-wide alerts without coupling UI components.")

heading3("GlobalErrorBoundary")
kv("React Error Boundary", "A custom React class component that catches JavaScript errors in child component trees and displays a fallback UI instead of crashing the entire application. Applied to the Browse Jobs and Job Detail routes.")

heading3("ProtectedRoute Component")
kv("Role-Based Access Control", "A custom React component that wraps protected routes. It checks the current user's role from AuthContext and either renders the route's component (if the role is allowed) or redirects to the appropriate login page. Used for /admin, /recruiter, /dashboard, /profile, and /saved-jobs routes.")

heading3("ScrollToTop Component")
kv("ScrollToTop", "A React utility component that scrolls the browser window to the top whenever the route changes — ensuring a consistent user experience when navigating between pages.")

heading3("UUID (Universally Unique Identifier)")
kv("uuid v9.0.1", "Every database document (Profile, Job, Application) is assigned a custom UUID v4 string id field in addition to MongoDB's default _id (ObjectId). This dual-ID approach allows the frontend to use either ID format when making API requests.")

heading3("Base64 File Storage Strategy")
kv("Strategy", "Since Render's free tier does not provide persistent disk storage, all uploaded files (resumes, identity documents, certificates, profile pictures) are converted to Base64-encoded data URLs in memory (using Multer's memoryStorage) and stored directly as string fields in the MongoDB documents. This approach requires no external file storage service (like AWS S3) and works within the free-tier infrastructure.")
kv("Trade-off", "Base64 encoding increases file size by approximately 33%. MongoDB document size limit is 16 MB, so the 10 MB file upload limit ensures documents stay within bounds.")

heading3("Tailwind CSS Configuration")
kv("tailwind.config.js", "Custom configuration file extending Tailwind's default theme. The content array points to all .jsx and .js files in /src so that Tailwind's JIT (Just-in-Time) compiler can purge unused CSS classes and produce a minimal production CSS bundle.")

# ─── 3.7 SYSTEM ARCHITECTURE ─────────────────────────────────────────────────
doc.add_page_break()
heading2("3.7  System Architecture Summary")
body(
    "The GlobalAKJobs platform follows a three-tier client-server architecture:"
)
kv("Tier 1 — Presentation Layer", "React SPA (hosted on Netlify/Vercel CDN). The user's browser loads the SPA and communicates with the backend via HTTPS REST API calls.")
kv("Tier 2 — Application Layer", "Express.js REST API (hosted on Render.com). All business logic, authentication, data validation, and file processing happens here.")
kv("Tier 3 — Data Layer", "MongoDB Atlas (cloud-hosted NoSQL database). Stores all persistent application data including users, jobs, applications, and documents.")

body("\nExternal Services integrated with the platform:")
kv("Google Cloud (OAuth 2.0)", "Provides Google Sign-In authentication for candidates and agents")
kv("Resend.com", "Transactional email delivery for password reset emails in production")
kv("Nodemailer + Ethereal", "Email capture and preview in local development")
kv("GitHub", "Source code repository and CI/CD trigger for automatic deployments on both Render and Netlify")

# ─── FOOTER ──────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(200)
r = p.add_run("— End of Document —")
r.italic = True
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
r.font.size = Pt(11)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sub.add_run("GlobalAKJobs  |  Maldives Career Platform  |  Technical Documentation  |  2026")
sr2.font.size = Pt(9)
sr2.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = r"c:\Users\RITHI\OneDrive\Desktop\Maldivess\GlobalAKJobs_Project_Documentation.docx"
doc.save(out_path)
print(f"✅ Document saved: {out_path}")
