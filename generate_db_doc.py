from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── PAGE MARGINS ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

TEAL   = RGBColor(0x0D, 0x94, 0x88)
NAVY   = RGBColor(0x0F, 0x17, 0x2A)
DKTEAL = RGBColor(0x13, 0x4E, 0x4A)
MUTED  = RGBColor(0x47, 0x55, 0x69)
LIGHT  = RGBColor(0x94, 0xA3, 0xB8)

# ── HELPERS ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TEAL
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.italic = True; r.font.size = Pt(11); r.font.color.rgb = DKTEAL
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); r.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + level*0.3)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11)
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    k = p.add_run(f"{key}: "); k.bold = True; k.font.size = Pt(11)
    v = p.add_run(value);      v.font.size = Pt(11)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '0D9488')
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def add_table(headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        # Style header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        # Background colour (teal)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0D9488')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            # Alternating row shading
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0FDFA')
                tcPr.append(shd)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacer
    return table

# =============================================================================
#  TITLE PAGE
# =============================================================================
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
tr = tp.add_run("GlobalAKJobs — Maldives Career Platform")
tr.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = TEAL

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("DATABASE DESIGN DOCUMENT").font.size = Pt(14)

doc.add_paragraph()
dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run("Database: GlobalAKJobs-DB  |  Platform: MongoDB Atlas  |  March 2026")
dr.font.size = Pt(11); dr.font.color.rgb = MUTED

doc.add_page_break()

# =============================================================================
#  SECTION 1 — OVERVIEW
# =============================================================================
h1("1. Database Design Overview")
divider()
body(
    "The GlobalAKJobs platform uses MongoDB — a NoSQL document-oriented database — hosted on "
    "MongoDB Atlas (cloud). The database is named GlobalAKJobs-DB and contains 11 collections. "
    "Each collection stores documents in BSON (Binary JSON) format, giving the flexibility of "
    "schema-less design while Mongoose ODM enforces a well-defined schema at the application layer."
)
body(
    "Although MongoDB is a non-relational database and does not enforce foreign-key constraints "
    "at the engine level, the GlobalAKJobs application maintains logical relationships between "
    "collections using reference fields (storing ObjectIds or custom UUID strings). This approach "
    "is the standard MongoDB pattern for referencing documents across collections — equivalent to "
    "foreign keys in a relational database."
)

h2("Database at a Glance")
add_table(
    ["Collection", "Mongoose Model", "Description", "Documents (Live)"],
    [
        ["profiles",      "Profile",      "All users: Candidates, Agents, Admins",            "Active"],
        ["jobs",          "Job",          "Published job listings",                           "Active"],
        ["applications",  "Application",  "Candidate job applications with documents",        "Active"],
        ["agencies",      "Agency",       "Recruitment agency records",                       "Active"],
        ["jobrequests",   "JobRequest",   "Agent job-posting requests pending admin approval","Active"],
        ["notifications", "Notification", "In-app alerts for users",                          "Active"],
        ["documents",     "Document",     "Standalone user documents (separate from apps)",   "Active"],
        ["subscriptions", "Subscription", "User job subscriptions / alerts",                  "Active"],
        ["categories",    "Category",     "Job category reference list",                      "Active"],
        ["uploads.files", "GridFS",       "GridFS file metadata (legacy upload feature)",     "Present"],
        ["uploads.chunks","GridFS",       "GridFS binary file chunks (legacy)",                "Present"],
    ],
    [1.3, 1.2, 2.7, 1.0]
)

# =============================================================================
#  SECTION 2 — DATABASE STRUCTURE (COLLECTION-BY-COLLECTION)
# =============================================================================
doc.add_page_break()
h1("2. Database Structure")
divider()
body(
    "This section documents every collection in GlobalAKJobs-DB with its full field list, "
    "data types, constraints, and purpose. Field definitions are sourced from both the live "
    "MongoDB Atlas schema (verified via MCP inspection) and the Mongoose model definitions in "
    "the application source code."
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.1  profiles Collection  (Mongoose Model: Profile)")
body(
    "Central user store for all three roles in the system: CANDIDATE (job seekers), "
    "AGENT (recruiters / agencies), and ADMIN (platform administrators). "
    "Every other collection references a Profile document via its id or _id field."
)
add_table(
    ["Field", "BSON Type", "Required", "Unique", "Description"],
    [
        ["_id",                   "ObjectId", "Auto",  "Yes", "MongoDB auto-generated primary key"],
        ["id",                    "String",   "Auto",  "Yes", "Custom UUID v4 — used as the public-facing user ID across all references"],
        ["role",                  "String",   "Yes",   "No",  "Enum: ADMIN | AGENT | CANDIDATE — determines access level"],
        ["full_name",             "String",   "Yes",   "No",  "User's full display name"],
        ["email",                 "String",   "Yes",   "Yes", "Unique login email address (stored lowercase)"],
        ["password",              "String",   "Yes",   "No",  "Account password (plain text in current build; bcrypt hashing recommended for production)"],
        ["contact_number",        "String",   "No",    "No",  "User's phone number"],
        ["agency_name",           "String",   "Agents","No",  "Required for AGENT accounts — name of the recruitment agency"],
        ["license_number",        "String",   "No",    "No",  "Agency license number (optional, for agents)"],
        ["skills",                "Array[]",  "No",    "No",  "Array of skill strings — relevant for CANDIDATE profiles"],
        ["experience_years",      "Number",   "No",    "No",  "Total years of work experience (default: 0)"],
        ["location",              "String",   "No",    "No",  "User's current location / home island"],
        ["avatar",                "String",   "No",    "No",  "Base64-encoded data URL of the profile photo (stored in DB, excluded from localStorage)"],
        ["savedJobs",             "Array[]",  "No",    "No",  "Array of Job id strings saved/bookmarked by a CANDIDATE"],
        ["status",                "String",   "No",    "No",  "Enum: ACTIVE | INACTIVE | BANNED | PENDING (default: ACTIVE; agents start as PENDING until approved)"],
        ["requiresPasswordChange","Boolean",  "No",    "No",  "True if admin created the agent with a temp password and first-login change is required (default: false)"],
        ["temporaryPassword",     "String",   "No",    "No",  "Admin-generated one-time password for agent first login; cleared after password change"],
        ["agencyId",              "ObjectId", "No",    "No",  "Reference to agencies._id — links agent profile to its agency record"],
        ["googleId",              "String",   "No",    "No",  "Google OAuth unique sub ID — used to link Google account to this profile; prevents duplicate accounts"],
        ["resetPasswordToken",    "String",   "No",    "No",  "Cryptographically-secure 32-byte hex token for password-reset email flow"],
        ["resetPasswordExpires",  "Date",     "No",    "No",  "Expiry timestamp for the reset token (15 minutes from generation)"],
        ["createdAt",             "Date",     "Auto",  "No",  "Mongoose timestamps — document creation time"],
        ["updatedAt",             "Date",     "Auto",  "No",  "Mongoose timestamps — last modification time"],
    ],
    [1.5, 1.0, 0.9, 0.7, 3.1]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.2  jobs Collection  (Mongoose Model: Job)")
body(
    "Stores all approved and published job listings on the platform. Jobs are created directly "
    "by Admins or auto-created when Admin approves a JobRequest submitted by an Agent."
)
add_table(
    ["Field", "BSON Type", "Required", "Description"],
    [
        ["_id",          "ObjectId","Auto", "MongoDB auto-generated primary key"],
        ["id",           "String",  "Auto", "Custom UUID v4 — public-facing job identifier"],
        ["title",        "String",  "Yes",  "Job title (e.g., 'Sous Chef', 'Civil Engineer')"],
        ["company",      "String",  "Yes",  "Hiring company or organisation name"],
        ["location",     "String",  "Yes",  "Job location / island (e.g., 'Malé', 'Addu City')"],
        ["category",     "String",  "No",   "Enum: Hospitality | Construction | Healthcare | IT | Education | Retail | Manufacturing | Tourism | Fishing | Agriculture | Other"],
        ["salary_range", "String",  "No",   "Salary range (e.g., '$800 – $1,200 / month')"],
        ["description",  "String",  "Yes",  "Full job description text"],
        ["requirements", "Array[]", "No",   "Array of requirement strings listed for the role"],
        ["vacancies",    "Number",  "No",   "Number of open positions (default: 1, from live DB schema)"],
        ["status",       "String",  "No",   "Enum: OPEN | CLOSED — controls public visibility"],
        ["posted_date",  "Date",    "Auto", "Date the job was posted (defaults to now)"],
        ["createdAt",    "Date",    "Auto", "Mongoose timestamps — creation time"],
        ["updatedAt",    "Date",    "Auto", "Mongoose timestamps — last update time"],
    ],
    [1.3, 1.0, 0.95, 3.95]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.3  applications Collection  (Mongoose Model: Application)")
body(
    "Records every job application submitted by a candidate (directly or via an agent). "
    "All supporting documents (resume, identity, certificates, PCC, good standing certificate) "
    "are embedded directly in this document as Base64-encoded strings — no external file server is needed."
)
add_table(
    ["Field", "BSON Type", "Required", "Description"],
    [
        ["_id",                      "ObjectId","Auto", "MongoDB auto-generated primary key"],
        ["id",                       "String",  "Auto", "Custom UUID v4 — public-facing application identifier"],
        ["job_id",                   "String",  "Yes",  "References jobs.id — the job being applied to"],
        ["candidate_name",           "String",  "Yes",  "Full name of the applicant"],
        ["email",                    "String",  "Yes",  "Applicant's email address"],
        ["contact_number",           "String",  "Yes",  "Applicant's contact number"],
        ["agent_id",                 "String",  "No",   "References profiles.id — set only when application is submitted by an agent on behalf of a candidate"],
        ["resume",                   "Object",  "No",   "Embedded object: { filename, contentType, data (Base64) } — applicant's CV"],
        ["identity",                 "Object",  "No",   "Embedded object: { filename, contentType, data (Base64) } — passport or national ID"],
        ["certificates",             "Object",  "No",   "Embedded object: { filename, contentType, data (Base64) } — educational / professional certificates"],
        ["pcc",                      "Object",  "No",   "Embedded object: { filename, contentType, data (Base64) } — Police Clearance Certificate"],
        ["goodStanding",             "Object",  "No",   "Embedded object: { filename, contentType, data (Base64) } — Good Standing Certificate"],
        ["status",                   "String",  "No",   "Enum: PENDING | REVIEWING | ACCEPTED | REJECTED | APPROVED | HOLD | SELECTED (default: PENDING)"],
        ["visibility_request_status","String",  "No",   "Enum: NOT_REQUESTED | PENDING | APPROVED | REJECTED — candidate request to view their application progress"],
        ["visibility_requested_at",  "Date",    "No",   "Timestamp when candidate requested visibility"],
        ["visibility_reviewed_by",   "String",  "No",   "ID / name of the agent or admin who reviewed the visibility request"],
        ["visibility_reviewed_at",   "Date",    "No",   "Timestamp when visibility request was reviewed"],
        ["applied_at",               "Date",    "Auto", "Timestamp of application submission (defaults to now)"],
        ["createdAt",                "Date",    "Auto", "Mongoose timestamps — creation time"],
        ["updatedAt",                "Date",    "Auto", "Mongoose timestamps — last update time"],
    ],
    [1.7, 1.0, 0.95, 3.55]
)

# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
h2("2.4  agencies Collection  (Mongoose Model: Agency)")
body(
    "Stores the details of recruitment agencies registered on the platform. "
    "Agent profiles (in the profiles collection) may reference an agency via agencyId."
)
add_table(
    ["Field", "BSON Type", "Required", "Unique", "Description"],
    [
        ["_id",        "ObjectId","Auto","Yes", "MongoDB auto-generated primary key (referenced by profiles.agencyId)"],
        ["name",       "String",  "Yes", "Yes", "Agency name — must be unique across the platform"],
        ["email",      "String",  "Yes", "No",  "Agency contact email"],
        ["contact",    "String",  "No",  "No",  "Agency telephone / mobile number"],
        ["location",   "String",  "No",  "No",  "Physical office location / island"],
        ["logo",       "String",  "No",  "No",  "URL to the agency logo image (default placeholder)"],
        ["status",     "String",  "No",  "No",  "Enum: Pending | Active | Rejected — admin-controlled approval status"],
        ["description","String",  "No",  "No",  "Short description of the agency and its services"],
        ["website",    "String",  "No",  "No",  "Agency website URL"],
        ["createdAt",  "Date",    "Auto","No",  "Document creation timestamp"],
    ],
    [1.2, 1.0, 0.9, 0.7, 3.4]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.5  jobrequests Collection  (Mongoose Model: JobRequest)")
body(
    "An intermediary collection in the job-posting workflow. When an Agent wants to post a job, "
    "they submit a JobRequest. An Admin reviews it and either approves (creating a Job document) "
    "or rejects it. This prevents agents from directly modifying the live jobs collection."
)
add_table(
    ["Field", "BSON Type", "Required", "Description"],
    [
        ["_id",             "ObjectId","Auto", "MongoDB auto-generated primary key"],
        ["id",              "String",  "Auto", "Custom UUID v4"],
        ["agent_id",        "String",  "Yes",  "References profiles.id — the agent who submitted the request"],
        ["agent_name",      "String",  "Yes",  "Agent's full name (denormalised for quick display)"],
        ["agent_email",     "String",  "Yes",  "Agent's email (denormalised for notifications)"],
        ["agency_name",     "String",  "No",   "Name of the submitting agency (denormalised)"],
        ["title",           "String",  "Yes",  "Proposed job title"],
        ["company",         "String",  "Yes",  "Hiring company name"],
        ["location",        "String",  "Yes",  "Job location"],
        ["category",        "String",  "Yes",  "Enum: same values as jobs.category"],
        ["salary_range",    "String",  "No",   "Proposed salary range"],
        ["description",     "String",  "Yes",  "Full job description"],
        ["requirements",    "Array[]", "No",   "List of job requirements"],
        ["vacancies",       "Number",  "No",   "Number of vacancies requested (default: 1)"],
        ["status",          "String",  "No",   "Enum: PENDING | APPROVED | REJECTED (default: PENDING)"],
        ["reviewed_by",     "String",  "No",   "ID or name of the Admin who reviewed the request"],
        ["review_notes",    "String",  "No",   "Admin's comments or reason for rejection"],
        ["reviewed_at",     "Date",    "No",   "Timestamp when Admin made the decision"],
        ["approved_job_id", "String",  "No",   "References jobs.id — set when request is approved and a Job document is created"],
        ["createdAt",       "Date",    "Auto", "Mongoose timestamps — creation time"],
        ["updatedAt",       "Date",    "Auto", "Mongoose timestamps — last update time"],
    ],
    [1.5, 1.0, 0.9, 3.8]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.6  notifications Collection  (Mongoose Model: Notification)")
body(
    "Stores in-app notifications sent to users when key events occur "
    "(e.g., application status changes, new job alerts, system messages)."
)
add_table(
    ["Field", "BSON Type", "Required", "Description"],
    [
        ["_id",                  "ObjectId","Auto", "MongoDB auto-generated primary key"],
        ["userId",               "String",  "Yes",  "References profiles.id — the recipient user (indexed for fast lookup)"],
        ["title",                "String",  "Yes",  "Short notification title"],
        ["message",              "String",  "Yes",  "Full notification body text"],
        ["type",                 "String",  "No",   "Enum: JOB_ALERT | APPLICATION_UPDATE | SYSTEM (default: SYSTEM)"],
        ["metadata.jobId",       "String",  "No",   "Optional: job ID related to the notification"],
        ["metadata.applicationId","String", "No",   "Optional: application ID related to the notification"],
        ["isRead",               "Boolean", "No",   "Whether the user has read/dismissed the notification (default: false)"],
        ["createdAt",            "Date",    "Auto", "Notification creation timestamp (default: now)"],
    ],
    [1.7, 1.0, 0.9, 3.6]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.7  documents Collection  (Mongoose Model: Document)")
body(
    "Stores standalone user-uploaded documents not tied to a specific job application "
    "(e.g., a candidate pre-uploading their resume or passport to their profile library)."
)
add_table(
    ["Field",          "BSON Type","Required","Description"],
    [
        ["_id",          "ObjectId","Auto","MongoDB auto-generated primary key"],
        ["user_id",      "String",  "Yes", "References profiles.id — the document owner (indexed for fast lookups)"],
        ["document_type","String",  "Yes", "Type label: e.g., RESUME | PASSPORT | LICENSE | CERTIFICATE"],
        ["filename",     "String",  "Yes", "Original filename of the uploaded file"],
        ["content_type", "String",  "Yes", "MIME type (e.g., application/pdf, image/jpeg)"],
        ["file_data",    "String",  "Yes", "Base64-encoded binary content of the file"],
        ["created_at",   "Date",    "Auto","Upload timestamp (default: now)"],
    ],
    [1.4, 1.0, 0.95, 3.85]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.8  subscriptions Collection  (Mongoose Model: Subscription)")
body(
    "Tracks which users have subscribed to job alerts for specific jobs. "
    "A compound unique index on (userId, jobId) prevents duplicate subscriptions."
)
add_table(
    ["Field",    "BSON Type","Required","Unique","Description"],
    [
        ["_id",      "ObjectId","Auto","Yes", "MongoDB auto-generated primary key"],
        ["userId",   "String",  "Yes", "No",  "References profiles.id — the subscribing user (indexed)"],
        ["jobId",    "String",  "Yes", "No",  "References jobs.id or jobs._id — the job being subscribed to"],
        ["createdAt","Date",    "Auto","No",  "Subscription creation timestamp (default: now)"],
    ],
    [1.2, 1.0, 0.9, 0.7, 3.4]
)
body("Compound Unique Index: { userId: 1, jobId: 1 } — ensures one subscription per user per job.")

# ─────────────────────────────────────────────────────────────────────────────
h2("2.9  categories Collection  (Mongoose Model: Category)")
body(
    "A simple reference list of job categories used for filtering and normalisation. "
    "The job categories are also enforced as an enum in the jobs and jobrequests models."
)
add_table(
    ["Field",    "BSON Type","Required","Description"],
    [
        ["_id",      "ObjectId","Auto","MongoDB auto-generated primary key"],
        ["name",     "String",  "Yes", "Category name (e.g., Hospitality, Construction, IT, Healthcare)"],
        ["createdAt","Date",    "Auto","Creation timestamp"],
    ],
    [1.3, 1.1, 0.9, 3.9]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("2.10  uploads.files & uploads.chunks  (GridFS)")
body(
    "These two collections are part of MongoDB's GridFS specification — a mechanism for storing "
    "large binary files (> 16 MB) by splitting them into chunks. They were created during an "
    "early version of the file upload feature. The current implementation uses Base64 embedding "
    "directly in the applications and documents collections (avoiding the need for GridFS). "
    "These collections remain present in the database for legacy compatibility."
)
add_table(
    ["Collection",    "Purpose"],
    [
        ["uploads.files",  "Stores file metadata: filename, length, chunkSize, uploadDate, contentType"],
        ["uploads.chunks", "Stores binary file data split into 255 KB chunks; references uploads.files._id"],
    ],
    [2.0, 5.2]
)

# =============================================================================
#  SECTION 3 — NORMALIZATION
# =============================================================================
doc.add_page_break()
h1("3. Database Normalization")
divider()
body(
    "Normalization is the process of organizing a database to reduce redundancy and improve data "
    "integrity. In relational databases, this is achieved through Normal Forms (1NF, 2NF, 3NF, BCNF). "
    "Although MongoDB is a NoSQL database and does not require strict normalization, the GlobalAKJobs "
    "schema has been designed following normalization principles where appropriate and uses strategic "
    "denormalization only when it improves performance."
)

h2("3.1  First Normal Form (1NF)")
body(
    "1NF requires that: (a) each column holds atomic (indivisible) values, (b) each column contains "
    "values of a single type, and (c) each row is uniquely identifiable."
)
bullet("✔ Every document in every collection has a unique identifier (_id ObjectId auto-assigned by MongoDB).")
bullet("✔ Scalar fields (full_name, email, title, status, etc.) hold a single atomic value.")
bullet("✔ Array fields (skills, requirements, savedJobs) satisfy 1NF because NoSQL documents permit multi-valued attributes as first-class arrays — this is the standard and recommended MongoDB pattern.")
bullet("✔ Embedded document fields (resume, identity, pcc, etc. within applications) group related sub-attributes together atomically — consistent with NoSQL 1NF compliance.")
bullet("✔ The profiles collection uses a single email field (unique index) as the natural key, in addition to the system _id.")

h2("3.2  Second Normal Form (2NF)")
body(
    "2NF requires 1NF and that all non-key attributes are fully functionally dependent on the entire "
    "primary key (no partial dependencies on composite keys)."
)
bullet("✔ MongoDB documents use a single _id as primary key — there are no composite primary keys, so partial dependency is structurally impossible.")
bullet("✔ All fields in each collection are attributes of the entity that collection represents:")
bullet("profiles — every field (name, email, skills, avatar, etc.) describes the user", level=1)
bullet("jobs — every field (title, company, salary, category) describes the job posting", level=1)
bullet("applications — every field describes one application event", level=1)
bullet("✔ 2NF is fully satisfied across all collections.")

h2("3.3  Third Normal Form (3NF)")
body(
    "3NF requires 2NF and that there are no transitive dependencies — non-key attributes must not "
    "depend on other non-key attributes."
)
bullet("✔ In the profiles collection, all fields depend directly on the user's identity (id/_id). There are no fields that depend on another non-key field.")
bullet("✔ In the jobs collection, all fields describe the job itself. Category is a string value (not a derived field), and its validity is enforced by the Mongoose enum constraint rather than by a separate join.")
bullet("✔ In the applications collection, candidate_name and email are stored directly with the application document. While these values also exist in profiles, they are intentionally denormalised here — this is a deliberate design choice explained in Section 3.5.")
bullet("✔ In jobrequests, agent_name, agent_email, and agency_name are denormalised from the profiles collection. This is intentional for performance reasons (see 3.5).")
bullet("ℹ For a strict relational 3NF interpretation: the jobrequests collection would be considered to have a transitive dependency (agency_name depends on agent_id → profiles → agency_name). This is an accepted NoSQL trade-off.")

h2("3.4  BCNF (Boyce-Codd Normal Form)")
body(
    "BCNF is a stronger form of 3NF that requires every determinant to be a candidate key. "
    "In the context of MongoDB:"
)
bullet("✔ Each collection is keyed by a single unique _id. Every other field is determined by that key only.")
bullet("✔ The profiles collection has two unique identifiers: _id (ObjectId) and id (UUID). Both are candidate keys. email is also a unique field and can be considered a natural key. These do not create BCNF violations because they are all equivalent identifiers for the same entity.")
bullet("✔ The subscriptions collection uses a compound unique index {userId, jobId} — this compound key uniquely determines the subscription, satisfying BCNF.")
bullet("✔ Overall BCNF is satisfied at the collection level.")

h2("3.5  Intentional Denormalization (NoSQL Design Pattern)")
body(
    "Denormalization — storing redundant data to improve read performance — is a standard and "
    "recommended practice in MongoDB. The following instances of denormalization are intentional "
    "and justified in GlobalAKJobs:"
)
add_table(
    ["Location", "Denormalized Field(s)", "Source Collection", "Justification"],
    [
        ["applications", "candidate_name, email, contact_number",
         "profiles",
         "Applications must remain self-contained historical records. If the candidate later updates their profile, the application document should still show what was submitted at the time of applying."],
        ["jobrequests", "agent_name, agent_email, agency_name",
         "profiles",
         "Admin dashboard displays request details without requiring a JOIN. Avoids extra API calls. Agent info is captured at time of submission."],
        ["notifications", "userId (UUID string)",
         "profiles",
         "Direct user ID reference avoids a lookup and enables indexed queries on notifications per user."],
        ["jobs", "category (enum string)",
         "categories",
         "Category is embedded as a string rather than a reference ID. This is intentional: job cards render instantly without a JOIN to the categories collection."],
    ],
    [1.2, 1.6, 1.2, 3.2]
)

# =============================================================================
#  SECTION 4 — RELATIONSHIPS
# =============================================================================
doc.add_page_break()
h1("4. Database Relationships")
divider()
body(
    "MongoDB does not enforce foreign-key relationships at the engine level. Instead, GlobalAKJobs "
    "maintains logical relationships by storing reference IDs (either MongoDB ObjectIds or custom "
    "UUID strings) in fields of one collection that point to documents in another. These references "
    "are resolved at the application layer (Express.js / Mongoose). The following subsections document "
    "every relationship in the database."
)

h2("4.1  Relationship Diagram (Entity-Relationship Overview)")
body(
    "The entities and their relationships are described below. In MongoDB terminology:\n"
    "• Reference (Ref) = logical foreign key stored as a string or ObjectId\n"
    "• Embed = sub-document stored directly inside the parent document (no separate collection)\n"
    "• 1:1 = One-to-One    • 1:N = One-to-Many    • N:N = Many-to-Many"
)

add_table(
    ["From Collection", "Field", "→ To Collection", "Type", "Nature"],
    [
        ["profiles",    "agencyId",         "agencies._id",        "Ref (ObjectId)", "Many:1 — Many agents can belong to one agency"],
        ["applications","job_id",           "jobs.id",             "Ref (UUID)",     "Many:1 — Many applications belong to one job"],
        ["applications","agent_id",         "profiles.id",         "Ref (UUID)",     "Many:1 — Many applications can be submitted by one agent"],
        ["applications","resume/identity/pcc/certificates/goodStanding", "—",        "Embed",          "1:1 — Documents embedded within the application document"],
        ["jobrequests", "agent_id",         "profiles.id",         "Ref (UUID)",     "Many:1 — Many requests can come from one agent"],
        ["jobrequests", "approved_job_id",  "jobs.id",             "Ref (UUID)",     "1:1 — One approved request creates one job"],
        ["notifications","userId",          "profiles.id",         "Ref (UUID)",     "Many:1 — Many notifications belong to one user"],
        ["subscriptions","userId",          "profiles.id",         "Ref (UUID)",     "Many:Many — A user can subscribe to many jobs; a job can have many subscribers"],
        ["subscriptions","jobId",           "jobs.id",             "Ref (UUID)",     "Many:Many — (compound unique index enforces one subscription per user-job pair)"],
        ["documents",   "user_id",          "profiles.id",         "Ref (UUID)",     "Many:1 — Many documents belong to one user"],
        ["profiles",    "savedJobs[]",      "jobs.id",             "Ref Array (UUID)","Many:Many — A candidate can save many jobs; a job can be saved by many candidates"],
    ],
    [1.4, 1.5, 1.5, 1.2, 2.6]
)

# ─────────────────────────────────────────────────────────────────────────────
h2("4.2  Detailed Relationship Descriptions")

h3("4.2.1  Profile ↔ Agency  (Many-to-One)")
body(
    "An Agent profile references an Agency document using the agencyId field (MongoDB ObjectId). "
    "Multiple agents can belong to the same agency, making this a Many-to-One relationship. "
    "The Agency document stores details like the agency name, email, website, and approval status. "
    "When the Admin views an agent's profile, they can use agencyId to fetch the corresponding "
    "agency document if needed."
)
kv("Field",       "profiles.agencyId  →  agencies._id")
kv("Cardinality", "Many Agents : One Agency")
kv("Enforcement", "Application layer — Mongoose populates agency details if needed")

h3("4.2.2  Application ↔ Job  (Many-to-One)")
body(
    "Each application document stores the job_id referencing the job it was submitted for. "
    "A single job can receive many applications from different candidates, making this a "
    "Many-to-One relationship. The agent or admin looking at applications for a job queries "
    "the applications collection by job_id."
)
kv("Field",       "applications.job_id  →  jobs.id  (UUID string)")
kv("Cardinality", "Many Applications : One Job")
kv("Query Pattern","db.applications.find({ job_id: '<uuid>' })")

h3("4.2.3  Application ↔ Agent Profile  (Many-to-One, Optional)")
body(
    "If a recruitment agent submits an application on behalf of a candidate, the agent_id field "
    "in the application document stores the Agent's UUID (from profiles.id). This field is optional "
    "— it is null / absent for candidate-submitted applications. This allows the recruiter dashboard "
    "to filter and show only applications submitted by the logged-in agent."
)
kv("Field",       "applications.agent_id  →  profiles.id  (UUID string, optional)")
kv("Cardinality", "Many Applications : One Agent (optional)")

h3("4.2.4  Application ↔ Uploaded Documents  (Embedded One-to-One)")
body(
    "Rather than storing uploaded files in a separate collection, the application document itself "
    "embeds five document sub-objects: resume, identity, certificates, pcc (Police Clearance "
    "Certificate), and goodStanding (Good Standing Certificate). Each sub-object has three fields: "
    "filename, contentType (MIME), and data (Base64 string). This embedded design means a single "
    "MongoDB read retrieves the entire application along with all its documents — no JOIN needed."
)
kv("Pattern",     "Embedded document (sub-object) — not a separate collection reference")
kv("Trade-off",   "Increased document size; mitigated by 10 MB upload limit (MongoDB max doc size is 16 MB)")

h3("4.2.5  JobRequest ↔ Agent Profile  (Many-to-One)")
body(
    "A job request is always submitted by an agent. The agent_id, agent_name, and agent_email "
    "fields in the jobrequests collection identify the submitting agent. The agent_id references "
    "profiles.id. The name and email are denormalised copies (captured at time of submission) "
    "so the admin dashboard can display them without a separate lookup."
)
kv("Field",       "jobrequests.agent_id  →  profiles.id  (UUID string)")
kv("Cardinality", "Many JobRequests : One Agent")

h3("4.2.6  JobRequest ↔ Job  (One-to-One, after Approval)")
body(
    "When an Admin approves a JobRequest, a new Job document is created and the resulting job's "
    "UUID is written back to the jobrequests.approved_job_id field. This creates a One-to-One "
    "traceability link: every approved job can be traced back to the request that generated it. "
    "REJECTED requests leave approved_job_id as null."
)
kv("Field",       "jobrequests.approved_job_id  →  jobs.id  (UUID string, null until approved)")
kv("Cardinality", "One JobRequest : One Job  (after approval)")

h3("4.2.7  Notification ↔ Profile  (Many-to-One)")
body(
    "Each notification targets one user, stored as userId (profiles.id UUID). An indexed query "
    "on userId allows the backend to instantly retrieve all notifications for a specific user. "
    "The metadata sub-object optionally embeds a jobId or applicationId to provide context "
    "for the notification without requiring a Mongoose populate."
)
kv("Field",       "notifications.userId  →  profiles.id  (UUID string, indexed)")
kv("Cardinality", "Many Notifications : One User")

h3("4.2.8  Subscription ↔ Profile & Job  (Many-to-Many)")
body(
    "The subscriptions collection acts as a junction table (associative entity) implementing "
    "a Many-to-Many relationship between users and jobs. A single user can subscribe to many "
    "jobs, and a single job can have many subscribers. A compound unique index on "
    "(userId, jobId) prevents duplicate subscriptions."
)
kv("Fields",      "subscriptions.userId  →  profiles.id  |  subscriptions.jobId  →  jobs.id")
kv("Cardinality", "Many-to-Many (Users ↔ Jobs via subscriptions)")
kv("Index",       "{ userId: 1, jobId: 1 }  — compound unique index")

h3("4.2.9  Document ↔ Profile  (Many-to-One)")
body(
    "The documents collection stores files uploaded by a user to their personal document library "
    "(independent of any specific application). The user_id field references profiles.id. "
    "One user can have many stored documents (resume, passport, license, etc.)."
)
kv("Field",       "documents.user_id  →  profiles.id  (UUID string, indexed)")
kv("Cardinality", "Many Documents : One User")

h3("4.2.10  Profile.savedJobs ↔ Jobs  (Many-to-Many, Embedded Array)")
body(
    "Candidates can bookmark/save jobs for later review. The saved job IDs are stored as a string "
    "array inside the Profile document (profiles.savedJobs[]). Each element references a jobs.id. "
    "This avoids the need for a separate junction collection for the save-jobs feature. The trade-off "
    "is that the array must be maintained carefully if a job is deleted."
)
kv("Field",       "profiles.savedJobs[]  →  jobs.id  (Array of UUID strings)")
kv("Cardinality", "Many-to-Many (Candidates save many Jobs; stored in profile array)")
kv("Alternative", "A subscriptions-style junction collection could be used instead for cleaner normalisation")

# =============================================================================
#  SECTION 5 — INDEXES
# =============================================================================
doc.add_page_break()
h1("5. Database Indexes")
divider()
body(
    "Indexes allow MongoDB to find documents efficiently without scanning every document in a collection. "
    "The following indexes are defined or used in GlobalAKJobs-DB."
)
add_table(
    ["Collection",    "Field(s)",           "Index Type",    "Purpose"],
    [
        ["profiles",    "_id",                "Default",       "MongoDB auto-index on primary key"],
        ["profiles",    "id (UUID)",          "Unique",        "Ensures no duplicate UUIDs; used by all API lookups"],
        ["profiles",    "email",              "Unique",        "Prevents duplicate email registrations; used for login lookup"],
        ["profiles",    "googleId",           "—",             "Looked up during Google OAuth to find existing accounts"],
        ["jobs",        "_id",                "Default",       "MongoDB auto-index on primary key"],
        ["jobs",        "id (UUID)",          "Unique",        "Public-facing job ID for application and detail page lookups"],
        ["applications","_id",               "Default",       "MongoDB auto-index on primary key"],
        ["applications","id (UUID)",         "Unique",        "Public-facing application ID"],
        ["applications","job_id",            "—",             "Queried frequently to list all applications for a job"],
        ["jobrequests", "id (UUID)",         "Unique",        "Public-facing request ID"],
        ["notifications","userId",           "Index",         "Enables fast retrieval of all notifications for a user"],
        ["documents",   "user_id",           "Index",         "Enables fast retrieval of all documents for a user"],
        ["subscriptions","userId + jobId",   "Compound Unique","Prevents duplicate user-job subscriptions"],
        ["agencies",    "name",              "Unique",        "Prevents two agencies registering with the same name"],
    ],
    [1.3, 1.5, 1.2, 3.2]
)

# =============================================================================
#  SECTION 6 — DATA INTEGRITY & CONSTRAINTS
# =============================================================================
h1("6. Data Integrity and Constraints")
divider()
body(
    "Because MongoDB does not natively enforce relational constraints (like foreign keys or cascades), "
    "data integrity in GlobalAKJobs is maintained through Mongoose schema validations and application-level logic."
)
add_table(
    ["Constraint Type", "Where Applied", "How Enforced"],
    [
        ["Required fields",    "All collections",      "Mongoose 'required: true' — rejects document save if field is missing"],
        ["Unique email",       "profiles",             "Mongoose unique index — returns duplicate key error on duplicate email"],
        ["Unique UUID id",     "profiles, jobs, applications, jobrequests", "Mongoose unique index on id field"],
        ["Enum validation",    "profiles.role, profiles.status, jobs.category, jobs.status, applications.status, jobrequests.status, notifications.type, agencies.status", "Mongoose enum — rejects values not in the allowed list"],
        ["Reference integrity","applications.job_id, applications.agent_id, jobrequests.agent_id, etc.", "Application-layer lookup — server returns 404 if referenced document not found"],
        ["File size limit",    "applications, documents", "Multer 10 MB limit enforced before file is parsed"],
        ["Token expiry",       "profiles.resetPasswordExpires", "Query filters resetPasswordExpires > now — expired tokens are rejected"],
        ["Compound unique",    "subscriptions",        "Mongoose compound unique index {userId, jobId}"],
        ["Google token verify","profiles.googleId",    "Backend verifies Google ID token via Google's tokeninfo API before creating/linking account"],
        ["Account status check","profiles.status",     "Login route checks status !== BANNED before allowing access; Google OAuth checks for ACTIVE status"],
    ],
    [1.6, 2.2, 3.4]
)

# =============================================================================
#  FOOTER
# =============================================================================
doc.add_page_break()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(200)
fr = fp.add_run("— End of Database Design Document —")
fr.italic = True; fr.font.color.rgb = LIGHT; fr.font.size = Pt(11)

sp2 = doc.add_paragraph()
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sp2.add_run("GlobalAKJobs  |  Database: GlobalAKJobs-DB  |  MongoDB Atlas  |  2026")
sr2.font.size = Pt(9); sr2.font.color.rgb = LIGHT

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r"c:\Users\RITHI\OneDrive\Desktop\Maldivess\GlobalAKJobs_Database_Design.docx"
doc.save(out)
print(f"✅ Saved: {out}")

