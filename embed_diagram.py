"""
Adds the ER diagram image into the existing Database Design Word document.
Inserts it right after Section 4 heading (before detailed relationship descriptions).
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc_path = r'c:\Users\RITHI\OneDrive\Desktop\Maldivess\GlobalAKJobs_Database_Design.docx'
img_path = r'c:\Users\RITHI\OneDrive\Desktop\Maldivess\er_diagram.png'
out_path = r'c:\Users\RITHI\OneDrive\Desktop\Maldivess\GlobalAKJobs_Database_Design.docx'

doc = Document(doc_path)

TEAL = RGBColor(0x0D, 0x94, 0x88)
NAVY = RGBColor(0x0F, 0x17, 0x2A)

# ── Find insertion point: paragraph containing "4.1  Relationship Diagram" ──
# We will replace the existing section 4.1 body text with our image + caption,
# then keep everything else. Simplest approach: find the paragraph, then insert
# after it.

target_para_idx = None
for i, para in enumerate(doc.paragraphs):
    if '4.1' in para.text and 'Relationship Diagram' in para.text:
        target_para_idx = i
        break

if target_para_idx is None:
    print("⚠ Could not find paragraph 4.1 — appending diagram at end instead")
    # Append at end
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ENTITY RELATIONSHIP DIAGRAM (ERD)")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEAL

    doc.add_picture(img_path, width=Inches(7.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph("Figure 1: GlobalAKJobs Entity Relationship Diagram — All 9 Collections and their Relationships")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
else:
    # ── Insert picture paragraph AFTER paragraph at target_para_idx ──────────
    # We need to work at the XML level to insert a paragraph at a specific position.

    # First build the new paragraphs we want to insert
    # 1. Caption heading
    # 2. The image paragraph
    # 3. Figure caption

    # Step 1: find the XML parent of the target paragraph
    target_para = doc.paragraphs[target_para_idx]
    body = doc.element.body

    # Helper: create a centred paragraph with text
    def make_text_para(text, bold=False, italic=False, size=11, color=None, center=False):
        from docx.oxml.ns import nsmap
        from lxml import etree
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        if center:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
        spAft = OxmlElement('w:spacing')
        spAft.set(qn('w:after'), '60')
        pPr.append(spAft)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size*2))); rPr.append(sz)
        sz2 = OxmlElement('w:szCs'); sz2.set(qn('w:val'), str(int(size*2))); rPr.append(sz2)
        if bold:
            b = OxmlElement('w:b'); rPr.append(b)
        if italic:
            it = OxmlElement('w:i'); rPr.append(it)
        if color:
            clr = OxmlElement('w:color'); clr.set(qn('w:val'), color); rPr.append(clr)
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = text; t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        return p

    # Reference node to insert after
    ref_node = target_para._p

    # Insert spacer paragraph
    sp = make_text_para('', size=6)
    ref_node.addnext(sp)
    ref_node = sp

    # Insert figure caption paragraph
    cap_p = make_text_para(
        'Figure 1: GlobalAKJobs Entity Relationship Diagram — '
        'All 9 Collections with Fields and Relationships',
        italic=True, size=9, color='64748B', center=True)
    ref_node.addnext(cap_p)
    ref_node = cap_p

    # Insert spacer
    sp2 = make_text_para('', size=4)
    ref_node.addnext(sp2)
    ref_node = sp2

    # Insert the image using a run with drawing element
    # Use docx's built-in add_picture approach on a temp doc, then move XML
    import copy
    tmp_doc = Document()
    tmp_doc.add_picture(img_path, width=Inches(7.0))
    pic_para = tmp_doc.paragraphs[-1]
    # Centre align
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_xml = copy.deepcopy(pic_para._p)
    ref_node.addnext(pic_xml)
    ref_node = pic_xml

    # Insert sub-heading before image
    hdr_p = make_text_para('Entity Relationship Diagram (ERD)',
                            bold=True, size=12, color='0D9488', center=True)
    ref_node_hdr = target_para._p
    # insert right after 4.1 paragraph
    # We already have the image chain; insert header right after target_para
    target_para._p.addnext(hdr_p)

    print(f"✅ Diagram inserted after paragraph: '{target_para.text[:60]}'")

doc.save(out_path)
print(f"✅ Updated document saved: {out_path}")
